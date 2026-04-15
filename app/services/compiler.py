from docx import Document
import io


def compile_book_docx(title: str, chapters: list):
    doc = Document()
    doc.add_heading(title, 0)

    for ch in chapters:
        doc.add_heading(f"Chapter {ch['chapter_number']}", level=1)
        doc.add_paragraph(ch['content'])
        doc.add_page_break()

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
